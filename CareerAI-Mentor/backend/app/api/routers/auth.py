from io import BytesIO
from typing import Annotated

import fitz
from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload
from app.core.security import decode_access_token
from app.db.database import get_db
from app.models import Resume, StudentProfile, User
from app.schemas import CurrentUserResponse, LoginRequest, OnboardingRequest, RegisterRequest, StudentProfileResponse, TokenResponse
from app.services.auth import AuthService

router = APIRouter(prefix="/api/auth", tags=["authentication"])
auth_service = AuthService(); security = HTTPBearer(auto_error=False)
DatabaseSession = Annotated[Session, Depends(get_db)]
def get_current_user(db: DatabaseSession, credentials: Annotated[HTTPAuthorizationCredentials | None, Depends(security)]) -> User:
    if credentials is None: raise HTTPException(status_code=401, detail="Not authenticated", headers={"WWW-Authenticate": "Bearer"})
    user = db.scalar(select(User).options(selectinload(User.profile)).where(User.id == decode_access_token(credentials.credentials)))
    if user is None: raise HTTPException(status_code=401, detail="User not found", headers={"WWW-Authenticate": "Bearer"})
    return user
CurrentUser = Annotated[User, Depends(get_current_user)]
@router.post("/register", response_model=TokenResponse, status_code=status.HTTP_201_CREATED)
def register(payload: RegisterRequest, db: DatabaseSession) -> TokenResponse: return TokenResponse(access_token=auth_service.register(db, payload))
@router.post("/login", response_model=TokenResponse)
def login(payload: LoginRequest, db: DatabaseSession) -> TokenResponse: return TokenResponse(access_token=auth_service.login(db, payload))
@router.get("/me", response_model=CurrentUserResponse)
def me(user: CurrentUser) -> User: return user
@router.post("/logout", status_code=status.HTTP_204_NO_CONTENT)
def logout(_: CurrentUser) -> None: return None

@router.get("/profile", response_model=StudentProfileResponse | None)
def get_profile(user: CurrentUser) -> StudentProfile | None:
    return user.profile

@router.put("/onboarding", response_model=StudentProfileResponse)
def save_onboarding(payload: OnboardingRequest, user: CurrentUser, db: DatabaseSession) -> StudentProfile:
    profile = user.profile or StudentProfile(user_id=user.id)
    for field, value in payload.model_dump().items():
        setattr(profile, field, value)
    db.add(profile)
    db.commit()
    db.refresh(profile)
    return profile

@router.post("/resume/upload")
def upload_resume(user: CurrentUser, db: DatabaseSession, file: UploadFile = File(...)) -> dict:
    if not file.filename:
        raise HTTPException(415, "Please choose a TXT, PDF, or DOCX resume file")
    extension = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if extension not in {"txt", "pdf", "docx"}:
        raise HTTPException(415, "Resume format must be TXT, PDF, or DOCX")
    raw = file.file.read()
    if not raw or len(raw) > 5 * 1024 * 1024:
        raise HTTPException(413, "Resume must be between 1 byte and 5 MB")
    try:
        if extension == "txt":
            text = raw.decode("utf-8", errors="replace")
        elif extension == "pdf":
            document = fitz.open(stream=raw, filetype="pdf")
            text = "\n".join(page.get_text() for page in document)
            document.close()
        else:
            document = Document(BytesIO(raw))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            text = "\n".join(paragraphs + table_text)
    except Exception as exc:
        raise HTTPException(422, "We could not read this resume file. Try exporting it again as PDF, DOCX, or TXT.") from exc
    if not text.strip():
        raise HTTPException(422, "No readable text was found in this resume. Please upload a text-based PDF, DOCX, or TXT file.")
    catalog = ["Python", "SQL", "JavaScript", "TypeScript", "React", "FastAPI", "Machine Learning", "Pandas", "NumPy", "Git", "Docker"]
    skills = [skill for skill in catalog if skill.lower() in text.lower()]
    score = min(100, 30 + 8 * len(skills) + 10 * sum(word in text.lower() for word in ["education", "project", "experience"]))
    resume = Resume(user_id=user.id, filename=file.filename, extracted_text=text, skills=skills, score=score)
    db.add(resume); db.commit(); db.refresh(resume)
    return {"id": resume.id, "filename": resume.filename, "skills": resume.skills, "score": resume.score}

@router.get("/resume")
def get_resumes(user: CurrentUser, db: DatabaseSession) -> list[dict]:
    rows = db.scalars(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc())).all()
    return [{"id": r.id, "filename": r.filename, "skills": r.skills, "score": r.score, "created_at": r.created_at} for r in rows]


ROLE_SKILLS = {
    "AI Engineer": ["Python", "Machine Learning", "FastAPI", "Docker", "Git"],
    "Python Developer": ["Python", "SQL", "FastAPI", "Git", "Docker"],
    "Data Analyst": ["Python", "SQL", "Pandas", "NumPy", "Data Visualization"],
    "Data Scientist": ["Python", "SQL", "Pandas", "Machine Learning", "Statistics"],
    "ML Engineer": ["Python", "Machine Learning", "Docker", "Git", "AWS"],
    "Full Stack Developer": ["JavaScript", "TypeScript", "React", "SQL", "Git"],
}

@router.post("/skill-gap/analyze")
def analyze_skill_gap(user: CurrentUser, db: DatabaseSession) -> dict:
    profile = user.profile
    if not profile or not profile.target_role:
        raise HTTPException(422, "Complete your student profile and choose a target role first")
    required = ROLE_SKILLS.get(profile.target_role)
    if not required:
        raise HTTPException(422, "Target role is not in the current role catalog")
    observed = {skill.casefold() for skill in profile.technical_skills}
    latest = db.scalar(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc()))
    if latest:
        observed.update(skill.casefold() for skill in latest.skills)
    matching = [skill for skill in required if skill.casefold() in observed]
    missing = [skill for skill in required if skill.casefold() not in observed]
    readiness = round(100 * len(matching) / len(required))
    return {"target_role": profile.target_role, "required_skills": required, "matching_skills": matching, "missing_skills": missing, "readiness_score": readiness, "method": "comparison of saved profile, latest uploaded resume, and seed role requirements"}

@router.post("/roadmap/generate")
def generate_roadmap(user: CurrentUser, db: DatabaseSession, weeks: int = 8) -> dict:
    if weeks not in {4, 8, 12, 16}:
        raise HTTPException(422, "Roadmap duration must be 4, 8, 12, or 16 weeks")
    profile = user.profile
    if not profile or not profile.target_role:
        raise HTTPException(422, "Complete onboarding first")
    required = ROLE_SKILLS.get(profile.target_role, [])
    known = {s.casefold() for s in profile.technical_skills}
    latest = db.scalar(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc()))
    if latest: known.update(s.casefold() for s in latest.skills)
    gaps = [s for s in required if s.casefold() not in known] or required[:2]
    tasks = []
    for week in range(1, weeks + 1):
        skill = gaps[(week - 1) % len(gaps)]
        tasks.append({"week": week, "goal": f"Build {skill} capability", "topics": [skill], "tasks": [f"Study core {skill} concepts", f"Complete focused {skill} practice"], "estimated_hours": profile.weekly_study_time or 5, "milestone": f"Week {week}: demonstrate one {skill} exercise"})
    return {"target_role": profile.target_role, "weeks": weeks, "based_on_missing_skills": gaps, "roadmap": tasks, "source": "deterministic plan from your profile, resume, and selected role"}

PRACTICE_QUESTIONS = [
    {"id": "python-list-1", "topic": "Python", "difficulty": "Easy", "type": "MCQ", "question": "What does len([1, 2, 3]) return?", "options": ["2", "3", "4"], "answer": "3", "explanation": "len returns the number of items."},
    {"id": "sql-having-1", "topic": "SQL", "difficulty": "Medium", "type": "MCQ", "question": "Which clause filters grouped SQL rows?", "options": ["WHERE", "HAVING", "ORDER BY"], "answer": "HAVING", "explanation": "HAVING filters after grouping and aggregation."},
]

@router.get("/practice/questions")
def practice_questions(user: CurrentUser) -> list[dict]:
    return [{k: v for k, v in question.items() if k != "answer"} | {"source": "seed/demo data"} for question in PRACTICE_QUESTIONS]

@router.post("/practice/{question_id}/submit")
def submit_practice(question_id: str, answer: str, user: CurrentUser) -> dict:
    question = next((q for q in PRACTICE_QUESTIONS if q["id"] == question_id), None)
    if not question:
        raise HTTPException(404, "Question not found")
    correct = answer.strip().casefold() == question["answer"].casefold()
    return {"correct": correct, "topic": question["topic"], "explanation": question["explanation"]}

@router.post("/interview/start")
def start_interview(user: CurrentUser, interview_type: str = "Technical") -> dict:
    if not user.profile or not user.profile.target_role:
        raise HTTPException(422, "Complete onboarding first")
    return {"type": interview_type, "target_role": user.profile.target_role, "question": f"Tell me about a project or skill that prepares you for a {user.profile.target_role} role."}

@router.post("/interview/evaluate")
def evaluate_interview(answer: str, user: CurrentUser) -> dict:
    words = len(answer.split())
    score = min(100, max(0, 35 + words * 2))
    feedback = "Add a specific project, your contribution, and a measurable outcome." if words < 20 else "Clear foundation. Strengthen it by naming a concrete outcome and technology choice."
    return {"overall_score": score, "communication_score": min(100, 40 + words * 2), "feedback": feedback, "basis": "response length and evidence-oriented answer guidance; not a personality assessment"}

@router.get("/recommendations/today")
def today_recommendations(user: CurrentUser, db: DatabaseSession) -> dict:
    profile = user.profile
    if not profile or not profile.target_role:
        raise HTTPException(422, "Complete onboarding first")
    required = ROLE_SKILLS.get(profile.target_role, [])
    observed = {s.casefold() for s in profile.technical_skills}
    latest = db.scalar(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc()))
    if latest: observed.update(s.casefold() for s in latest.skills)
    missing = [s for s in required if s.casefold() not in observed]
    actions = [{"priority": i + 1, "action": f"Study {skill}", "minutes": 30} for i, skill in enumerate(missing[:2])]
    actions.append({"priority": len(actions) + 1, "action": "Complete one focused practice question", "minutes": 20})
    return {"actions": actions[:3], "based_on": {"target_role": profile.target_role, "missing_skills": missing}}

@router.get("/dashboard/summary")
def dashboard_summary(user: CurrentUser, db: DatabaseSession) -> dict:
    profile = user.profile
    if not profile:
        return {"profile_complete": False, "readiness_score": None, "skills_recorded": 0, "resume_uploaded": False}
    required = ROLE_SKILLS.get(profile.target_role or "", [])
    observed = {s.casefold() for s in profile.technical_skills}
    latest = db.scalar(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc()))
    if latest: observed.update(s.casefold() for s in latest.skills)
    matching = [s for s in required if s.casefold() in observed]
    readiness = round(100 * len(matching) / len(required)) if required else None
    return {"profile_complete": bool(profile.target_role), "target_role": profile.target_role, "skills_recorded": len(observed), "resume_uploaded": latest is not None, "resume_score": latest.score if latest else None, "readiness_score": readiness, "missing_skills": [s for s in required if s not in matching]}

@router.post("/job-description/analyze")
def analyze_job_description(job_description: str, user: CurrentUser, db: DatabaseSession) -> dict:
    if len(job_description.strip()) < 30:
        raise HTTPException(422, "Provide a job description with at least 30 characters")
    catalog = sorted({skill for skills in ROLE_SKILLS.values() for skill in skills})
    required = [skill for skill in catalog if skill.casefold() in job_description.casefold()]
    observed = set(s.casefold() for s in (user.profile.technical_skills if user.profile else []))
    latest = db.scalar(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc()))
    if latest: observed.update(s.casefold() for s in latest.skills)
    matching = [skill for skill in required if skill.casefold() in observed]
    missing = [skill for skill in required if skill.casefold() not in observed]
    score = round(100 * len(matching) / len(required)) if required else 0
    return {"detected_required_skills": required, "matching_skills": matching, "missing_skills": missing, "job_match_score": score, "method": "deterministic matching against the known role-skill catalog"}

@router.get("/projects/recommendations")
def project_recommendations(user: CurrentUser, db: DatabaseSession) -> dict:
    profile = user.profile
    if not profile or not profile.target_role:
        raise HTTPException(422, "Complete onboarding first")
    required = ROLE_SKILLS.get(profile.target_role, [])
    known = {s.casefold() for s in profile.technical_skills}
    latest = db.scalar(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc()))
    if latest: known.update(s.casefold() for s in latest.skills)
    gaps = [s for s in required if s.casefold() not in known]
    focus = gaps[:3] or required[:3]
    projects = [
        {"name": f"{profile.target_role} Skill Tracker", "difficulty": "Beginner", "skills_learned": focus, "problem_statement": "Track learning tasks and demonstrate role-relevant fundamentals.", "steps": ["Design data model", "Build a focused feature", "Write a README and screenshots"]},
        {"name": f"{profile.target_role} Portfolio Project", "difficulty": "Intermediate", "skills_learned": required[:4], "problem_statement": "Build an end-to-end project that demonstrates the target role's core workflow.", "steps": ["Define user problem", "Implement core workflow", "Deploy or record a demo"]},
        {"name": "Interview Evidence Notebook", "difficulty": "Beginner", "skills_learned": ["Git", "Documentation", "Communication"], "problem_statement": "Document decisions, challenges, and measurable outcomes from your projects.", "steps": ["Create project templates", "Record weekly progress", "Prepare STAR interview stories"]},
    ]
    return {"target_role": profile.target_role, "based_on_missing_skills": gaps, "projects": projects, "source": "deterministic recommendations based on profile, resume, and role-skill catalog"}

@router.get("/feedback/summary")
def feedback_summary(user: CurrentUser, db: DatabaseSession) -> dict:
    profile = user.profile
    if not profile or not profile.target_role:
        raise HTTPException(422, "Complete onboarding first")
    required = ROLE_SKILLS.get(profile.target_role, [])
    known = {s.casefold() for s in profile.technical_skills}
    latest = db.scalar(select(Resume).where(Resume.user_id == user.id).order_by(Resume.created_at.desc()))
    if latest: known.update(s.casefold() for s in latest.skills)
    strengths = [s for s in required if s.casefold() in known]
    weak = [s for s in required if s.casefold() not in known]
    score = round(100 * len(strengths) / len(required)) if required else 0
    return {"overall_readiness": score, "strong_areas": strengths, "weak_areas": weak, "immediate_actions": [f"Study {skill}" for skill in weak[:2]] + ["Complete one practice question"], "next_milestone": "Demonstrate one project feature using a current weak skill", "evidence": {"profile_skills": profile.technical_skills, "latest_resume_uploaded": latest is not None, "resume_score": latest.score if latest else None}}
